from docx import Document
from docx.shared import Mm, Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START

from Laboratory import *
from Customer import *


def generate_protocol(lab, customer, date_off, date_izm):
    doc = Document()
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENTATION.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    current_section.left_margin = Cm(2)
    current_section.right_margin = Cm(2)
    current_section.top_margin = Cm(3)
    current_section.bottom_margin = Mm(15)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    # table0 = current_section.footer.add_table(rows=1, cols=2, width=50)
    # table0.cell(0, 0).text = f"Частичное или полное воспроизведение протокола запрещены без письменного разрешения" \
    #                          f" руководителя испытательной лаборатории Результаты исследований (испытаний), измерений " \
    #                          f"относятся только к объектам (образцам), прошедшим испытания, отбор" \
    #                          f"{customer.get_number_protocol()} от {date_off}"
    table1 = doc.add_table(rows=1, cols=2)
    table1_cells = table1.rows[0].cells
    table1_cells[0].paragraphs[0].add_run().add_picture('logo.jpg', width=Mm(75))
    p = table1_cells[1].paragraphs[0]
    r = p.add_run(lab.get_head())
    r.font.size = Pt(9)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table1.columns[0].width = Mm(75)
    table1_cells[0].width = Mm(75)
    table1.columns[1].width = Mm(165)
    table1_cells[1].width = Mm(165)
    table1.style = 'Medium List 1'
    doc.add_paragraph("")
    table2 = doc.add_table(rows=1, cols=2)
    table2_cells = table2.rows[0].cells
    p = table2_cells[1].paragraphs[0]
    p.add_run(lab.get_signature(date_off))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table2.columns[0].width = Mm(170)
    table2_cells[0].width = Mm(170)
    table2.columns[1].width = Mm(70)
    table2_cells[1].width = Mm(70)
    p = doc.add_paragraph()
    run = p.add_run(customer.get_number_protocol())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.bold = True
    p = doc.add_paragraph()
    customer.fill_text(date_izm, p)
    doc.add_page_break()
    lab.fill_measure([0, 1, 2], doc)
    # doc.add_paragraph("11. Сведения о средствах измерения:")
    # table3 = doc.add_table(rows=2, cols=6)
    # table3.cell(0, 0).merge(table3.cell(1, 0)).text = "Наименование средства измерения"
    # table3.cell(0, 1).merge(table3.cell(1, 1)).text = "Заводской номер"
    # table3.cell(0, 2).merge(table3.cell(0, 4)).text = "Свидетельство о государственной поверке"
    # table3.cell(1, 2).text = "Номер"
    # table3.cell(1, 3).text = "Выдано"
    # table3.cell(1, 4).text = "Действительно до"
    # table3.cell(0, 5).merge(table3.cell(1, 5)).text = "Погрешность измерения"
    # for i in range(0, 3):
    #     cells = table3.add_row().cells
    #     for j, item in enumerate(lab.get_measure(i)):
    #         cells[j].text = str(item)
    doc.add_paragraph("12. Идентификация используемого метода/методик (нормативно-техническая документация), "
                      "а также дополнительная информация, востребованная заказчиком (НД, необходимые для оценки):  ")
    table4 = doc.add_table(rows=1, cols=2)
    table4.cell(0, 0).text = "Область действия"
    table4.cell(0, 1).text = "Наименование нормативного документа"
    for i in range(0, 2):
        cells = table4.add_row().cells
        for j, item in enumerate(lab.get_methodology(i)):
            cells[j].text = str(item)
    doc.add_paragraph("13. Условия проведения исследований, испытаний (измерений), отбора образцов:")
    table5 = doc.add_table(rows=1, cols=5)
    table5.cell(0, 0).text = "№\nточки"
    table5.cell(0, 1).text = "Место измерений\n(наименование образца испытаний)"
    table5.cell(0, 2).text = "Температура\nвоздуха, oC"
    table5.cell(0, 3).text = "Атмосферное\nдавление,\nмм рт.ст."
    table5.cell(0, 4).text = "Относительная\nвлажность, %"
    customer.fill_weather_table(table5)
    doc.add_paragraph("14.  Результаты проверки работоспособности: уровни звукового давления на частотах калибратора, "
                      "полученные в конце измерений, отличаются от полученных в начале измерений не более чем на "
                      "0,5 дБА").bold = True
    doc.add_paragraph("15.  Временная характеристика шума: непостоянный, колеблющийся во времени;")
    doc.add_paragraph("16.  Результаты измерений параметров шума, дополнительная информация, востребованная "
                      "заказчиком:")
    table6 = doc.add_table(rows=2, cols=10)
    table6.cell(0, 0).merge(table6.cell(1, 0)).text = "№\nточки\n(рабочего\nместа)"
    table6.cell(0, 1).merge(table6.cell(1, 1)).text = "Место измерений\n(наименование образца\nиспытаний)*"
    table6.cell(0, 2).merge(table6.cell(1, 2)).text = "Источник шума**"
    table6.cell(0, 3).merge(table6.cell(1, 3)).text = "Характер шума**"
    table6.cell(0, 4).merge(table6.cell(0, 5)).text = "Уровень звука, дБА"
    table6.cell(1, 4).text = "Фактические\nзначения"
    table6.cell(1, 5).text = "Допустимые\nзначения"
    table6.cell(0, 6).merge(table6.cell(0, 7)).text = "Максимальный уровень звука,\nдБА"
    table6.cell(1, 6).text = "Результат\nизмерений"
    table6.cell(1, 7).text = "Нормативное\nзначение"
    table6.cell(0, 8).merge(table6.cell(0, 9)).text = "Эквивалентный уровень звука,\nдБА"
    table6.cell(1, 8).text = "Результат\nизмерений"
    table6.cell(1, 9).text = "Нормативное\nзначение"
    customer.fill_noise_table(table6)
    doc.add_paragraph("* Испытания проводились по месту осуществления деятельности Заказчика. В случае проведения "
                      "испытаний вне места осуществления деятельности Заказчика указывается адрес производственной "
                      "площадки.")
    doc.add_paragraph("** Указанные сведения предоставлены Заказчиком. Испытательная лаборатория не несет "
                      "ответственность за достоверность сведений, предоставленных Заказчиком.")
    doc.add_paragraph("17. Мнения и интерпретации: отсутствуют")
    doc.add_paragraph("18.  Дополнения, отклонения или исключения из метода: отсутствуют")
    doc.add_paragraph("19.  Дополнительная информация, востребованная заказчиком:")
    doc.add_paragraph(customer.get_hazard_wp())
    doc.add_paragraph("20. Измерения провел:")
    table7 = doc.add_table(rows=2, cols=3)
    lab.fill_signature(table7, 0)
    doc.add_paragraph("21. Протокол оформил")
    table8 = doc.add_table(rows=2, cols=3)
    lab.fill_signature(table8, 1)
    doc.save('test.docx')

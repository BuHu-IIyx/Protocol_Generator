from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Cm, Pt
from docx.oxml.shared import OxmlElement, qn


class Laboratory:
    def __init__(self, laboratory_id, short_name, name, name_lab, logo, director, address, certificate_number, phone,
                 e_mail):
        self.laboratory_id = laboratory_id
        self.name = name
        self.short_name = short_name
        self.name_lab = name_lab
        self.logo = logo
        self.director = director
        self.address = address
        self.certificate_number = certificate_number
        self.phone = phone
        self.e_mail = e_mail
        self.experts = []
        self.measuring = []
        self.methodologies = []

    class Experts:
        def __init__(self, name, position, certificate_number, status):
            self.name = name
            self.position = position
            self.certificate_number = certificate_number
            self.status = status

        def get_expert(self):
            return self.name + " " + self.position + " " + self.certificate_number

    class Measuring:
        def __init__(self, factory_number, name, accurate, verification_certificates):
            self.factory_number = factory_number
            self.name = name
            self.accurate = accurate
            self.verification_certificates = verification_certificates

        def get_measuring(self):
            return [self.name, self.factory_number, self.verification_certificates.get('number'),
                    self.verification_certificates.get('date_start'),
                    self.verification_certificates.get('date_off'), self.accurate]

        # class VerificationCertificate:
        #     def __init__(self, number, date_start, date_off):
        #         self.date_off = date_off
        #         self.date_start = date_start
        #         self.number = number
        #
        #     def get_verification(self):
        #         return [self.number, self.date_start, self.date_off]

        # def add_verification_certificate(self, number, date_start, date_off):
        #     new_certificate = self.VerificationCertificate(number, date_start, date_off)
        #     if len(self.verification_certificates) == 0:
        #         self.verification_certificates = [new_certificate]
        #     else:
        #         self.verification_certificates.append(new_certificate)

    class Methodology:
        def __init__(self, application, name):
            self.name = name
            self.application = application

        def get_meth(self):
            return [self.application, self.name]

    def add_expert(self, name, position, certificate_number, status):
        new_expert = self.Experts(name, position, certificate_number, status)
        if len(self.experts) == 0:
            self.experts = [new_expert]
        else:
            self.experts.append(new_expert)

    def add_measuring(self, factory_number, name, accurate, cert):
        new_measuring = self.Measuring(factory_number, name, accurate, cert)
        if len(self.measuring) == 0:
            self.measuring = [new_measuring]
        else:
            self.measuring.append(new_measuring)

    def add_methodology(self, application, name):
        new_methodology = self.Methodology(application, name)
        if len(self.methodologies) == 0:
            self.methodologies = [new_methodology]
        else:
            self.methodologies.append(new_methodology)

    def get_head(self):
        head_string = f"{self.name}\n{self.name_lab}\n{self.address}\n{self.phone}; {self.e_mail}\nУникальный номер " \
                      f"записи об аккредитации в реестре аккредитованных лиц: {self.certificate_number}\n"
        return head_string

    def get_signature(self, date):
        signature_string = f"УТВЕРЖДАЮ\nРуководитель испытательной лаборатории\n__________________{self.director}\n" \
                           f"{date}"
        return signature_string

    def get_measure(self, number_measure):
        return self.measuring[number_measure].get_measuring()

    def get_methodology(self, number_methodology):
        return self.methodologies[number_methodology].get_meth()

    def fill_signature(self, doc, status, person):
        status_list = ["20. Измерения провел:", "21. Протокол оформил"]
        doc.add_paragraph().add_run(status_list[status]).bold = True
        table = doc.add_table(rows=2, cols=3)
        for i in [0, 2]:
            cell = table.cell(0, i)
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), '4')
            cell._tc.get_or_add_tcPr().append(bottom_border)
        table.cell(0, 0).text = self.experts[person].position
        table.cell(1, 0).text = "(должность)"
        table.cell(0, 2).text = self.experts[person].name
        table.cell(1, 2).text = "(Ф.И.О.)"
        for i in range(0, 2):
            for j in range(0, 3):
                cell = table.cell(i, j)
                cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def fill_measure(self, doc):
        doc.add_paragraph().add_run("11. Сведения о средствах измерения:").bold = True
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        table_cells = table.rows[0].cells
        table.columns[0].width = Mm(80)
        table_cells[0].width = Mm(80)
        table.columns[1].width = Mm(40)
        table_cells[1].width = Mm(40)
        table.columns[2].width = Mm(15)
        table_cells[2].width = Mm(15)
        table.columns[3].width = Mm(15)
        table_cells[3].width = Mm(15)
        table.columns[4].width = Mm(15)
        table_cells[4].width = Mm(15)
        table.columns[5].width = Mm(80)
        table_cells[5].width = Mm(80)
        table.cell(0, 0).merge(table.cell(1, 0)).paragraphs[0].add_run("Наименование средства измерения")\
            .font.size = Pt(9)
        table.cell(0, 1).merge(table.cell(1, 1)).paragraphs[0].add_run("Заводской номер").font.size = Pt(9)
        table.cell(0, 2).merge(table.cell(0, 4)).paragraphs[0].add_run("Свидетельство о государственной поверке")\
            .font.size = Pt(9)
        table.cell(1, 2).paragraphs[0].add_run("Номер").font.size = Pt(9)
        table.cell(1, 3).paragraphs[0].add_run("Выдано").font.size = Pt(9)
        table.cell(1, 4).paragraphs[0].add_run("Действительно до").font.size = Pt(9)
        table.cell(0, 5).merge(table.cell(1, 5)).paragraphs[0].add_run("Погрешность измерения").font.size = Pt(9)

        for i in range(0, 2):
            for j in range(0, 6):
                cell = table.cell(i, j)
                cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                shading_elm = parse_xml(r'<w:shd {} w:fill="DEEAF6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

        for i in self.measuring:
            cells = table.add_row().cells
            for j, item in enumerate(i.get_measuring()):
                cells[j].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[j].paragraphs[0].add_run(str(item)).font.size = Pt(9)
        return 0

    def fill_methodology(self, doc):
        doc.add_paragraph().add_run("12. Идентификация используемого метода/методик (нормативно-техническая "
                                    "документация), а также дополнительная информация, востребованная заказчиком "
                                    "(НД, необходимые для оценки):").bold = True
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).paragraphs[0].add_run("Область действия").font.size = Pt(9)
        table.cell(0, 1).paragraphs[0].add_run("Наименование нормативного документа").font.size = Pt(9)
        table_cells = table.rows[0].cells
        table.columns[0].width = Mm(70)
        table_cells[0].width = Mm(70)
        table.columns[1].width = Mm(170)
        table_cells[1].width = Mm(170)
        for i in range(0, 2):
            cell = table.cell(0, i)
            cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading_elm = parse_xml(r'<w:shd {} w:fill="DEEAF6"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for i in self.methodologies:
            cells = table.add_row().cells
            for j, item in enumerate(i.get_meth()):
                cells[j].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[j].paragraphs[0].add_run(str(item)).font.size = Pt(9)


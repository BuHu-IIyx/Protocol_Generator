from docx.enum.section import WD_ORIENTATION
from docx import Document
from docx.oxml import ns
from Model.Laboratory import *
from Model.Customer import *


def generate_protocol_func(lab, customer, date_off, date_izm):
    # Create and customization document
    doc = Document()
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    current_section.orientation = WD_ORIENTATION.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height
    current_section.left_margin = Cm(2)
    current_section.right_margin = Cm(2)
    current_section.top_margin = Cm(3)
    current_section.bottom_margin = Mm(15)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    # Add 1st table in document
    table1 = doc.add_table(rows=1, cols=2)
    table1_cells = table1.rows[0].cells
    table1_cells[0].paragraphs[0].add_run().add_picture('logo.jpg', width=Mm(75))
    p = table1_cells[1].paragraphs[0]
    r = p.add_run(lab.get_head())
    r.font.size = Pt(9)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table1.columns[0].width = Mm(75)
    table1_cells[0].width = Mm(75)
    table1.columns[1].width = Mm(165)
    table1_cells[1].width = Mm(165)
    table1.style = 'Medium List 1'
    doc.add_paragraph("")
    table2 = doc.add_table(rows=1, cols=2)
    table2_cells = table2.rows[0].cells
    p = table2_cells[1].paragraphs[0]
    p.add_run(lab.get_signature(date_off))
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table2.columns[0].width = Mm(170)
    table2_cells[0].width = Mm(170)
    table2.columns[1].width = Mm(70)
    table2_cells[1].width = Mm(70)
    p = doc.add_paragraph()

    run = p.add_run(customer.get_number_protocol())
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run.bold = True
    p = doc.add_paragraph()
    customer.fill_text(date_izm, p)
    doc.add_page_break()
    lab.fill_measure(doc)
    doc.add_paragraph()
    lab.fill_methodology(doc)
    doc.add_paragraph()
    customer.fill_weather_table(doc)
    doc.add_paragraph().add_run("14.  Результаты проверки работоспособности: уровни звукового давления на частотах "
                                "калибратора, полученные в конце измерений, отличаются от полученных в начале измерений"
                                " не более чем на 0,5 дБА").bold = True
    doc.add_paragraph().add_run("15.  Временная характеристика шума: непостоянный, колеблющийся во времени;").\
        bold = True
    customer.fill_noise_table(doc)
    p = doc.add_paragraph()

    p.add_run("* Испытания проводились по месту осуществления деятельности Заказчика. В случае проведения "
              "испытаний вне места осуществления деятельности Заказчика указывается адрес производственной "
              "площадки.\n").font.size = Pt(9)
    p.add_run("** Указанные сведения предоставлены Заказчиком. Испытательная лаборатория не несет ответственность за "
              "достоверность сведений, предоставленных Заказчиком.").font.size = Pt(9)
    p = doc.add_paragraph()
    p.add_run("17. Мнения и интерпретации: ").bold = True
    p.add_run("отсутствуют.")
    p = doc.add_paragraph()
    p.add_run("18.  Дополнения, отклонения или исключения из метода: ").bold = True
    p.add_run("отсутствуют.")
    p = doc.add_paragraph()
    p.add_run("19.  Дополнительная информация, востребованная заказчиком:").bold = True
    p.add_run(customer.get_hazard_wp())
    lab.fill_signature(doc, 0, 0)
    lab.fill_signature(doc, 1, 1)
    # doc.sections[0].footer.paragraphs[0].text = "TEST FOOTER"
    table0 = doc.sections[0].footer.add_table(rows=1, cols=2, width=Cm(24))
    table0_cells = table0.rows[0].cells
    table0.columns[0].width = Mm(170)
    table0_cells[0].width = Mm(170)
    table0.columns[1].width = Mm(70)
    table0_cells[1].width = Mm(70)
    run = table0.cell(0, 0).paragraphs[0].add_run(f"Частичное или полное воспроизведение протокола запрещены без письменного разрешения" \
                             f" руководителя испытательной лаборатории Результаты исследований (испытаний), измерений " \
                             f"относятся только к объектам (образцам), прошедшим испытания, отбор " \
                             f"{customer.get_footer_number(date_off)}")
    run.italic = True
    run.font.size = Pt(8)
    add_page_number(table0.cell(0, 1).paragraphs[0])
    file_name = f'Output/{customer.contract_number}.docx'
    doc.save(file_name)


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    # выравниваем параграф
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    # запускаем динамическое обновление параграфа
    paragraph.add_run("Страница ")
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fld_char_1 = create_element('w:fldChar')
    create_attribute(fld_char_1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUMPAGES)
    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = "PAGE"
    # обозначаем конец позиции вывода
    fld_char_2 = create_element('w:fldChar')
    create_attribute(fld_char_2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fld_char_1)
    page_num_run._r.append(instr_text)
    page_num_run._r.append(fld_char_2)
    paragraph.add_run(" из ")
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fld_char_1 = create_element('w:fldChar')
    create_attribute(fld_char_1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUMPAGES)
    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = "NUMPAGES"
    # обозначаем конец позиции вывода
    fld_char_2 = create_element('w:fldChar')
    create_attribute(fld_char_2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fld_char_1)
    page_num_run._r.append(instr_text)
    page_num_run._r.append(fld_char_2)

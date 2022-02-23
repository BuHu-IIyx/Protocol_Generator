from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Mm, Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

class Customer:
    def __init__(self, name, legal_address, actual_address, contract_number, contract_date):
        self.contract_date = contract_date
        self.contract_number = contract_number
        self.actual_address = actual_address
        self.legal_address = legal_address
        self.name = name
        self.departments = []

    class Department:
        def __init__(self, name):
            self.name = name
            self.working_areas = []

        class WorkingArea:
            def __init__(self, number, name):
                self.name = name
                self.number = number
                self.hazard = False
                self.noise_params = 0
                self.weather_conditions = 0

            class NoiseParameter:
                def __init__(self, noise_source, nature_of_noise, sound_lvl, max_sound_lvl, eq_sound_lvl):
                    self.noise_source = noise_source
                    self.nature_of_noise = nature_of_noise
                    self.sound_lvl = sound_lvl
                    self.max_sound_lvl = max_sound_lvl
                    self.eq_sound_lvl = eq_sound_lvl

            class WeatherCondition:
                def __init__(self, temperature, atmo_pressure, humidity):
                    self.temperature = temperature
                    self.atmo_pressure = atmo_pressure
                    self.humidity = humidity

            def noise_parameter_ini(self, noise_source, nature_of_noise, sound_lvl, max_sound_lvl, eq_sound_lvl):
                self.noise_params = self.NoiseParameter(noise_source, nature_of_noise, sound_lvl, max_sound_lvl, eq_sound_lvl)
                if sound_lvl > 80 or max_sound_lvl > 110:
                    self.hazard = True

            def weather_condition_ini(self, temperature, atmo_pressure, humidity):
                self.weather_conditions = self.WeatherCondition(temperature, atmo_pressure, humidity)

            def get_noise_parameter(self):
                return [self.number, self.name, self.noise_params.noise_source, self.noise_params.nature_of_noise,
                        self.noise_params.sound_lvl, 80, self.noise_params.max_sound_lvl, 110,
                        self.noise_params.eq_sound_lvl, 80]

            def get_weather_condition(self):
                return [self.number, self.name, self.weather_conditions.temperature,
                        self.weather_conditions.atmo_pressure, self.weather_conditions.humidity]

        def add_working_area(self, number, name):
            new_working_area = self.WorkingArea(number, name)
            if len(self.working_areas) == 0:
                self.working_areas = [new_working_area]
            else:
                self.working_areas.append(new_working_area)

    def add_department(self, name):
        new_department = self.Department(name)
        if len(self.departments) == 0:
            self.departments = [new_department]
        else:
            self.departments.append(new_department)

    def get_hazard_wp(self):
        result = ""
        count = 0
        for department in self.departments:
            for area in department.working_areas:
                if area.hazard:
                    count += 1
                    result += area.number + ". " + area.name + "; "
        if count == 1:
            return f"по результатам измерений установлено: параметры шума на рабочем месте {result}не соответствуют " \
                   f"требованиям СанПиН 1.2.3685-21 «Гигиенические нормативы и требования к обеспечению безопасности " \
                   f"и (или) безвредности для человека факторов среды обитания». "
        elif count > 1:
            return f"по результатам измерений установлено: параметры шума на рабочих местах: {result}не соответствуют " \
                   f"требованиям СанПиН 1.2.3685-21 «Гигиенические нормативы и требования к обеспечению безопасности " \
                   f"и (или) безвредности для человека факторов среды обитания». "

        else:
            return "по результатам измерений установлено: параметры шума на рабочих местах соответствуют требованиям " \
                   "СанПиН 1.2.3685-21 «Гигиенические нормативы и требования к обеспечению безопасности и (или) " \
                   "безвредности для человека факторов среды обитания». "

    def get_number_protocol(self):
        return f"ПРОТОКОЛ № {self.contract_number}/Ш\nпроведения исследований, испытаний (измерений) шума"

    def get_footer_number(self, date):
        return f"Протокол №{self.contract_number}/Ш от {date}"

    def fill_text(self, date_izm, paragraph):
        paragraph.add_run("1. Наименование организации (заказчика): ").bold = True
        paragraph.add_run(f"{self.name}.\n")
        paragraph.add_run("2. Контактные данные заказчика (юридический адрес, фактический адрес места осуществления"
                          " деятельности): ").bold = True
        paragraph.add_run(f"{self.legal_address}.\n")
        paragraph.add_run("3. Место проведения испытаний (измерений): ").bold = True
        paragraph.add_run(f"{self.actual_address}.\n")
        paragraph.add_run("4. Дата осуществления лабораторной деятельности: ").bold = True
        paragraph.add_run(f"{date_izm}.\n")
        paragraph.add_run("5. Цель проведения измерений: ").bold = True
        paragraph.add_run("производственный контроль за соблюдением санитарно-эпидемиологических "
                          "требований и выполнением санитарно-противоэпидемических (профилактических) мероприятий.\n")
        paragraph.add_run("6. Описание и однозначная идентификация объекта (объектов) испытаний, состояние объекта "
                          "испытаний (при необходимости): ").bold = True
        paragraph.add_run("рабочие места на территории заказчика в соответствии с заявкой проведение "
                          "производственного контроля.\n")
        paragraph.add_run("7. Дата получения образца для испытаний: ").bold = True
        paragraph.add_run(f"в соответствии с заявкой на проведение измерений № {self.contract_number} от "
                          f"{self.contract_date}.\n")
        paragraph.add_run("8. Дата отбора образца: ").bold = True
        paragraph.add_run("отбор образцов не предусмотрен методикой измерений.\n")
        paragraph.add_run("9. Место отбора образцов: ").bold = True
        paragraph.add_run("отбор образцов не предусмотрен методикой измерений.\n")
        paragraph.add_run("10. Ссылка на план и методы отбора проб, используемые лабораторией или другими органами,"
                          " если они имеют отношение к достоверности и применению результатов: ").bold = True
        paragraph.add_run("в соответствии с журналом измерений.\n")

    def get_text(self, date_izm):
        p1 = f"1. Наименование организации (заказчика): {self.name}.\n"
        p2 = f"2. Контактные данные заказчика (юридический адрес, фактический адрес места осуществления деятельности):" \
             f" {self.legal_address}.\n"
        p3 = f"3. Место проведения испытаний (измерений): {self.actual_address}.\n"
        p4 = f"4. Дата осуществления лабораторной деятельности: {date_izm}.\n"
        p5 = "5. Цель проведения измерений: производственный контроль за соблюдением санитарно-эпидемиологических " \
             "требований и выполнением санитарно-противоэпидемических (профилактических) мероприятий.\n"
        p6 = "6. Описание и однозначная идентификация объекта (объектов) испытаний, состояние объекта испытаний (при " \
             "необходимости): рабочие места на территории заказчика в соответствии с заявкой проведение " \
             "производственного контроля.\n"
        p7 = f"7. Дата получения образца для испытаний: соответствии с заявкой на проведение измерений № " \
             f"{self.contract_number} от {self.contract_date}.\n"
        p8 = "8. Дата отбора образца: отбор образцов не предусмотрен методикой измерений.\n"
        p9 = "9.   Место отбора образцов: отбор образцов не предусмотрен методикой измерений.\n"
        p10 = "10. Ссылка на план и методы отбора проб, используемые лабораторией или другими органами, если они " \
              "имеют отношение к достоверности и применению результатов: в соответствии с журналом измерений.\n"
        return p1 + p2 + p3 + p4 + p5 + p6 + p7 + p8 + p9 + p10

    def fill_weather_table(self, doc):
        doc.add_paragraph().add_run("13. Условия проведения исследований, испытаний (измерений), отбора образцов:").\
            bold = True
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table_cells = table.rows[0].cells
        table.columns[0].width = Mm(10)
        table_cells[0].width = Mm(10)
        table.columns[1].width = Mm(170)
        table_cells[1].width = Mm(170)
        table.columns[2].width = Mm(20)
        table_cells[2].width = Mm(20)
        table.columns[3].width = Mm(20)
        table_cells[3].width = Mm(20)
        table.columns[4].width = Mm(20)
        table_cells[4].width = Mm(20)
        table.cell(0, 0).paragraphs[0].add_run("№\nточки").font.size = Pt(9)
        table.cell(0, 1).paragraphs[0].add_run("Место измерений\n(наименование образца испытаний)").font.size = Pt(9)
        table.cell(0, 2).paragraphs[0].add_run("Температура\nвоздуха, oC").font.size = Pt(9)
        table.cell(0, 3).paragraphs[0].add_run("Атмосферное\nдавление,\nмм рт.ст.").font.size = Pt(9)
        table.cell(0, 4).paragraphs[0].add_run("Относительная\nвлажность, %").font.size = Pt(9)
        for i in range(0, 5):
            cell = table.cell(0, i)
            cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shading_elm = parse_xml(r'<w:shd {} w:fill="DEEAF6"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for dept in self.departments:
            cells = table.add_row().cells
            cells[0].merge(cells[4]).paragraphs[0].add_run(dept.name).font.size = Pt(9)
            cells[0].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for area in dept.working_areas:
                area_cells = table.add_row().cells
                for i, item in enumerate(area.get_weather_condition()):
                    area_cells[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    area_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    area_cells[i].paragraphs[0].add_run(str(item)).font.size = Pt(9)
        return 0

    def fill_noise_table(self, doc):
        doc.add_paragraph().add_run(
            "16.  Результаты измерений параметров шума, дополнительная информация, востребованная "
            "заказчиком:").bold = True
        table = doc.add_table(rows=2, cols=10)
        table.style = 'Table Grid'
        table.cell(0, 0).merge(table.cell(1, 0)).paragraphs[0].add_run("№\nточки\n(рабочего\nместа)")\
            .font.size = Pt(9)
        table.cell(0, 1).merge(table.cell(1, 1)).paragraphs[0].add_run("Место измерений\n(наименование образца\n"
                                                                         "испытаний)*").font.size = Pt(9)
        table.cell(0, 2).merge(table.cell(1, 2)).paragraphs[0].add_run("Источник шума**").font.size = Pt(9)
        table.cell(0, 3).merge(table.cell(1, 3)).paragraphs[0].add_run("Характер шума**").font.size = Pt(9)
        table.cell(0, 4).merge(table.cell(0, 5)).paragraphs[0].add_run("Уровень звука, дБА").font.size = Pt(9)
        table.cell(1, 4).paragraphs[0].add_run("Фактические\nзначения").font.size = Pt(9)
        table.cell(1, 5).paragraphs[0].add_run("Допустимые\nзначения").font.size = Pt(9)
        table.cell(0, 6).merge(table.cell(0, 7)).paragraphs[0].add_run("Максимальный уровень звука, дБА")\
            .font.size = Pt(9)
        table.cell(1, 6).paragraphs[0].add_run("Результат\nизмерений").font.size = Pt(9)
        table.cell(1, 7).paragraphs[0].add_run("Нормативное\nзначение").font.size = Pt(9)
        table.cell(0, 8).merge(table.cell(0, 9)).paragraphs[0].add_run("Эквивалентный уровень звука, дБА")\
            .font.size = Pt(9)
        table.cell(1, 8).paragraphs[0].add_run("Результат\nизмерений").font.size = Pt(9)
        table.cell(1, 9).paragraphs[0].add_run("Нормативное\nзначение").font.size = Pt(9)
        for i in range(0, 2):
            for j in range(0, 10):
                cell = table.cell(i, j)
                cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                shading_elm = parse_xml(r'<w:shd {} w:fill="DEEAF6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

        for dept in self.departments:
            cells = table.add_row().cells
            cells[0].merge(cells[9]).paragraphs[0].add_run(dept.name).font.size = Pt(9)
            cells[0].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for area in dept.working_areas:
                area_cells = table.add_row().cells
                for i, item in enumerate(area.get_noise_parameter()):
                    area_cells[i].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    area_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    area_cells[i].paragraphs[0].add_run(str(item)).font.size = Pt(9)
        return 0


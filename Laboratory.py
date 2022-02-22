from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm, Cm, Pt


class Laboratory:
    def __init__(self, name, name_lab, logo, director, address, certificate_number, phone, e_mail):
        self.name = name
        self.name_lab = name_lab
        self.logo = logo
        self.director = director
        self.address = address
        self.certificate_number = certificate_number
        self.phone = phone
        self.e_mail = e_mail
        self.experts = []
        self.measuring = []
        self.methodologies = []

    class Experts:
        def __init__(self, name, position, certificate_number):
            self.name = name
            self.position = position
            self.certificate_number = certificate_number

        def get_expert(self):
            return self.name + " " + self.position + " " + self.certificate_number

    class Measuring:
        def __init__(self, factory_number, name, accurate):
            self.factory_number = factory_number
            self.name = name
            self.accurate = accurate
            self.verification_certificates = []

        def get_measuring(self):
            last_cert = len(self.verification_certificates) - 1
            return [self.name, self.factory_number, self.verification_certificates[last_cert].number,
                    self.verification_certificates[last_cert].date_start,
                    self.verification_certificates[last_cert].date_off, self.accurate]

        class VerificationCertificate:
            def __init__(self, number, date_start, date_off):
                self.date_off = date_off
                self.date_start = date_start
                self.number = number

            def get_verification(self):
                return [self.number, self.date_start, self.date_off]

        def add_verification_certificate(self, number, date_start, date_off):
            new_certificate = self.VerificationCertificate(number, date_start, date_off)
            if len(self.verification_certificates) == 0:
                self.verification_certificates = [new_certificate]
            else:
                self.verification_certificates.append(new_certificate)

    class Methodology:
        def __init__(self, application, name):
            self.name = name
            self.application = application

        def get_meth(self):
            return [self.application, self.name]

    def add_expert(self, name, position, certificate_number):
        new_expert = self.Experts(name, position, certificate_number)
        if len(self.experts) == 0:
            self.experts = [new_expert]
        else:
            self.experts.append(new_expert)

    def add_measuring(self, factory_number, name, accurate):
        new_measuring = self.Measuring(factory_number, name, accurate)
        if len(self.measuring) == 0:
            self.measuring = [new_measuring]
        else:
            self.measuring.append(new_measuring)

    def add_methodology(self, application, name):
        new_methodology = self.Methodology(application, name)
        if len(self.methodologies) == 0:
            self.methodologies = [new_methodology]
        else:
            self.methodologies.append(new_methodology)

    # def get_all(self):
    #     lab_string = "Название ООО: " + self.name + "\n" + "Название лаборатории: " + self.name_lab + "\n" + \
    #                  "Директор: " + self.director + "\n " + "Адрес: " + self.address + "\n" + \
    #                  "Уникальный номер записи об аккредитации в реестре аккредитованных лиц: " + \
    #                  self.certificate_number + "\nТелефон: " + self.phone + "\nE-mail: " + self.e_mail + "\n"
    #     lab_exp = ""
    #     for i in self.experts:
    #         lab_exp += i.get_expert() + "\n"
    #
    #     lab_meas = ""
    #     for i in self.measuring:
    #         lab_meas += i.get_measuring() + "\n"
    #
    #     lab_meth = ""
    #     for i in self.methodologies:
    #         lab_meth += i.get_meth() + "\n"
    #
    #     return lab_string + lab_exp + lab_meas + lab_meth

    def get_head(self):
        head_string = f"{self.name}\n{self.name_lab}\n{self.address}\n{self.phone}; {self.e_mail}\nУникальный номер " \
                      f"записи об аккредитации в реестре аккредитованных лиц: {self.certificate_number}\n"
        return head_string

    def get_signature(self, date):
        signature_string = f"УТВЕРЖДАЮ\nРуководитель испытательной лаборатории\n__________________{self.director}\n" \
                           f"{date}"
        return signature_string

    def get_measure(self, number_measure):
        return self.measuring[number_measure].get_measuring()

    def get_methodology(self, number_methodology):
        return self.methodologies[number_methodology].get_meth()

    def fill_signature(self, table, person):
        table.cell(0, 0).text = self.experts[person].position
        table.cell(1, 0).text = "(должность)"
        table.cell(0, 2).text = self.experts[person].name
        table.cell(1, 2).text = "(Ф.И.О.)"

    def fill_measure(self, measure_list, doc):
        doc.add_paragraph().add_run("11. Сведения о средствах измерения:").bold = True
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'
        table.cell(0, 0).merge(table.cell(1, 0)).paragraphs[0].add_run("Наименование средства измерения").font.size = Pt(9)
        table.cell(0, 1).merge(table.cell(1, 1)).paragraphs[0].add_run("Заводской номер").font.size = Pt(9)
        table.cell(0, 2).merge(table.cell(0, 4)).paragraphs[0].add_run("Свидетельство о государственной поверке").font.size = Pt(9)
        table.cell(1, 2).paragraphs[0].add_run("Номер").font.size = Pt(9)
        table.cell(1, 3).paragraphs[0].add_run("Выдано").font.size = Pt(9)
        table.cell(1, 4).paragraphs[0].add_run("Действительно до").font.size = Pt(9)
        table.cell(0, 5).merge(table.cell(1, 5)).paragraphs[0].add_run("Погрешность измерения").font.size = Pt(9)
        for i in range(0, 2):
            for j in range(0, 6):
                cell = table.cell(i, j)
                cell.paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                shading_elm = parse_xml(r'<w:shd {} w:fill="DEEAF6"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)

        for i in measure_list:
            cells = table.add_row().cells
            for j, item in enumerate(self.get_measure(i)):
                cells[j].paragraphs[0].alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[j].paragraphs[0].add_run(str(item)).font.size = Pt(9)
        return 0
